from __future__ import annotations

import csv
import math
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)


SOURCES = {
    "Ministerio de Sanidad - Nuevo Estatuto Marco": "https://www.sanidad.gob.es/areas/profesionesSanitarias/nuevoEstatutoMarco.htm",
    "Anteproyecto de Ley Estatuto Marco 08/01/2026": "https://www.sanidad.gob.es/areas/profesionesSanitarias/docs/APL_personal_estatutario_08_01_2026.pdf",
    "CESM - Comité de Huelga 08/04/2026": "https://www.cesm.org/2026/04/08/el-comite-de-huelga-denuncia-que-el-ministerio-entorpece-la-negociacion-para-esconder-su-incapacidad-para-llegar-a-acuerdos-que-atiendan-las-reivindicaciones-del-colectivo/",
    "CESM - Convocatoria huelga 22/01/2026": "https://www.cesm.org/2026/01/22/el-comite-de-huelga-convoca-un-paro-indefinido-a-partir-del-proximo-16-de-febrero-para-conseguir-un-estatuto-medico-y-facultativo/",
    "SATSE - Acuerdo con Sanidad 26/01/2026": "https://www.satse.es/notas-prensa/-/v/85301/acuerdo-inicio-tramitacion-estatuto-marco",
    "SIMEG - Retribuciones MIR 2026": "https://simeg.org/wp/wp-content/uploads/2026/04/Retribuciones-MIR-2026.pdf",
    "Aragon SALUD - Tablas 2026": "https://www.aragon.es/documents/20127/128260043/Tablas_retributivas_personal_estatutario_SALUD_2026.pdf/e527d18c-861e-0a7d-166e-cd07bf2aad77?download=true&t=1771839687505&version=1.0",
    "SAS Andalucia - Retribuciones 2026": "https://www.sspa.juntadeandalucia.es/servicioandaluzdesalud/profesionales/guia-laboral/retribuciones",
    "SCS Canarias - Retribuciones 2026": "https://www3.gobiernodecanarias.org/sanidad/scs/contenidoGenerico.jsp?idCarpeta=c32fdf85-fc15-11dd-a72f-93771b0e33f6&idDocument=b2540aca-fdab-11dd-b2aa-596c3deb1b4e",
    "SES Extremadura - Tablas 2026": "https://www.sindicato-sip.es/wp-content/uploads/2026/02/Tablas-retributivas-2026-SES.pdf",
    "BOCM Madrid - Retribuciones 2026": "https://www.bocm.es/boletin/CM_Orden_BOCM/2026/02/09/BOCM-20260209-6.PDF",
}


# Brutos mensuales sin guardias y valores brutos de guardia, compilados del informe SIMEG 2026
# y contrastados puntualmente con tablas oficiales autonómicas localizadas.
ROWS = [
    ("Andalucia", "SIMEG + SAS", [1380, 1490, 1628, 1766, 1904], [(14.07, 15.78, 28.14), (15.42, 17.28, 30.84), (18.02, 20.17, 36.04), (20.22, 22.61, 40.44)]),
    ("Aragon", "SIMEG + SALUD", [1387, 1498, 1637, 1776, 1914], [(17.34, 18.58, 34.68), (18.79, 20.13, 37.58), (20.23, 21.68, 40.46), (21.68, 23.23, 43.36)]),
    ("Asturias", "SIMEG", [1387, 1498, 1637, 1776, 1914], [(17.01, 18.18, 25.81), (19.32, 20.66, 29.32), (21.65, 23.14, 32.84), (23.97, 25.63, 36.36)]),
    ("Baleares", "SIMEG + Ibsalut", [1638, 1776, 1915, 2054, 2193], [(19.29, 21.21, 42.44), (21.08, 23.18, 46.36), (23.49, 25.83, 51.68), (25.91, 28.28, 56.99)]),
    ("Canarias", "SIMEG + SCS", [1591, 1702, 1841, 1979, 2118], [(11.20, 17.92, 22.40), (13.66, 21.86, 27.32), (16.15, 25.84, 32.30), (18.59, 29.74, 37.18)]),
    ("Cantabria", "SIMEG", [1380, 1489, 1626, 1763, 1899], [(16.80, 19.52, 31.43), (19.02, 20.78, 33.46), (21.47, 23.41, 37.68), (23.95, 26.07, 41.33)]),
    ("Castilla y Leon", "SIMEG", [1380, 1490, 1628, 1766, 1904], [(15.48, 17.34, 27.98), (17.74, 19.80, 31.90), (20.08, 22.38, 36.00), (22.34, 24.88, 40.10)]),
    ("Castilla-La Mancha", "SIMEG + SESCAM", [1387, 1498, 1637, 1776, 1914], [(18.21, 19.53, 26.34), (21.67, 23.23, 30.12), (24.28, 25.97, 33.85), (27.74, 29.70, 39.52)]),
    ("Cataluna", "SIMEG", [1504, 1618, 1757, 1897, 2044], [(18.88, 20.60, 41.17), (22.66, 24.71, 49.38), (26.42, 28.83, 57.63), (28.30, 30.86, 61.40)]),
    ("Com. Valenciana", "SIMEG + CESM-CV/SATSE", [1502, 1622, 1770, 1917, 2065], [(13.90, 15.90, 27.80), (16.42, 18.41, 32.84), (18.95, 20.95, 37.90), (21.48, 23.48, 42.96)]),
    ("Extremadura", "SIMEG + SES", [1414, 1525, 1664, 1803, 1941], [(15.16, 16.58, 21.96), (16.56, 17.94, 23.33), (17.94, 19.85, 24.71), (19.85, 21.26, 26.08)]),
    ("Galicia", "SIMEG", [1387, 1498, 1636, 1774, 1912], [(16.05, 17.18, 17.18), (18.37, 19.69, 19.69), (20.70, 22.17, 22.17), (22.90, 24.53, 24.53)]),
    ("La Rioja", "SIMEG", [1387, 1498, 1637, 1776, 1914], [(15.71, 16.47, 32.94), (18.57, 19.46, 38.92), (20.00, 20.96, 41.92), (22.85, 23.94, 47.88)]),
    ("Madrid", "SIMEG + BOCM", [1526, 1637, 1775, 1914, 2053], [(12.50, 14.87, 29.74), (15.00, 17.36, 34.72), (17.59, 19.85, 39.70), (19.99, 22.33, 44.66)]),
    ("Murcia", "SIMEG", [1474, 1498, 1637, 1776, 1914], [(16.56, 17.75, 33.12), (19.00, 21.39, 38.00), (21.06, 23.46, 42.12), (22.40, 25.95, 44.80)]),
    ("Navarra", "SIMEG", [1387, 1498, 1637, 1776, 1914], [(19.39, 19.39, 38.78), (22.60, 22.60, 45.20), (25.83, 25.83, 51.66), (28.45, 28.45, 56.90)]),
    ("Pais Vasco", "SIMEG", [1484, 1597, 1738, 1879, 2020], [(17.18, 18.90, 34.36), (18.75, 20.62, 37.50), (21.86, 24.06, 43.72), (24.99, 28.74, 49.98)]),
]


# Tabla con guardias de la imagen de CSIF/Pilar CTO, coincidente con el escenario SIMEG de 80 h/mes.
WITH_GUARDS = {
    "Andalucia": [(1925, 2557), (2070, 2780), (2296, 3134), (2484, 3456), (2563, 3594)],
    "Aragon": [(2083, 2796), (2232, 3024), (2386, 3280), (2533, 3536), (2611, 3675)],
    "Asturias": [(2089, 2783), (2285, 3084), (2484, 3414), (2677, 3743), (2757, 3881)],
    "Baleares": [(2407, 3239), (2575, 3526), (2771, 3865), (2963, 4198), (3041, 4337)],
    "Canarias": [(2012, 2689), (2243, 3041), (2469, 3423), (2691, 3801), (2770, 3940)],
    "Cantabria": [(2067, 2770), (2243, 3041), (2442, 3377), (2640, 3715), (2717, 3851)],
    "Castilla y Leon": [(2003, 2674), (2197, 2971), (2398, 3304), (2589, 3630), (2667, 3768)],
    "Castilla-La Mancha": [(2141, 2884), (2419, 3279), (2641, 3630), (2840, 4054), (2919, 4192)],
    "Cataluna": [(2279, 3066), (2542, 3492), (2810, 3943), (2982, 4237), (3065, 4385)],
    "Com. Valenciana": [(2000, 2674), (2210, 2995), (2418, 3346), (2624, 3696), (2708, 3843)],
    "Extremadura": [(2000, 2670), (2144, 2891), (2309, 3156), (2471, 3433), (2549, 3572)],
    "Galicia": [(2014, 2705), (2211, 3007), (2404, 3336), (2590, 3655), (2667, 3793)],
    "La Rioja": [(1998, 2667), (2223, 3011), (2375, 3266), (2592, 3636), (2671, 3775)],
    "Madrid": [(1948, 2597), (2151, 2907), (2361, 3250), (2556, 3583), (2635, 3722)],
    "Murcia": [(2107, 2834), (2274, 3090), (2452, 3394), (2615, 3674), (2694, 3813)],
    "Navarra": [(2192, 2938), (2404, 3306), (2615, 3703), (2861, 4052), (2875, 4190)],
    "Pais Vasco": [(2216, 2910), (2401, 3153), (2634, 3553), (2918, 3991), (2979, 4131)],
}

DISPLAY_NAMES = {
    "Andalucia": "Andalucía",
    "Aragon": "Aragón",
    "Cataluna": "Cataluña",
    "Castilla y Leon": "Castilla y León",
    "Pais Vasco": "País Vasco",
}

# Fila "Media de España" publicada en la tabla CSIF/Pilar CTO/SIMEG, con guardias.
PUBLISHED_WITH_GUARDS_MEAN = [
    (2087, 2806),
    (2284, 3114),
    (2486, 3452),
    (2684, 3790),
    (2758, 3929),
]


def euro(value: float, decimals: int = 0) -> str:
    s = f"{value:,.{decimals}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def display_ccaa(ccaa: str) -> str:
    return DISPLAY_NAMES.get(ccaa, ccaa)


ORDINARY_WEEKLY_HOURS = 35
ORDINARY_HOURS = ORDINARY_WEEKLY_HOURS * 52 / 12
GUARD_HOURS = 80
TOTAL_HOURS = ORDINARY_HOURS + GUARD_HOURS
TOTAL_HOURS_PLUS_FORMATION_5H = TOTAL_HOURS + (5 * 52 / 12)
ORDINARY_HOURS_37_5 = 37.5 * 52 / 12
TOTAL_HOURS_37_5 = ORDINARY_HOURS_37_5 + GUARD_HOURS


def means_no_guards() -> list[float]:
    return [sum(row[2][i] for row in ROWS) / len(ROWS) for i in range(5)]


def means_with_guards() -> list[tuple[float, float]]:
    return PUBLISHED_WITH_GUARDS_MEAN


def write_csv() -> Path:
    path = OUT / "datos_retribuciones_mir_2026.csv"
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow([
            "CCAA", "Fuente", "R1 sin guardias bruto mes", "R2 sin guardias bruto mes",
            "R3 sin guardias bruto mes", "R4 sin guardias bruto mes", "R5 sin guardias bruto mes",
            "R1 guardia lab €/h", "R1 guardia fest €/h", "R1 guardia especial €/h",
            "R2 guardia lab €/h", "R2 guardia fest €/h", "R2 guardia especial €/h",
            "R3 guardia lab €/h", "R3 guardia fest €/h", "R3 guardia especial €/h",
            "R4/R5 guardia lab €/h", "R4/R5 guardia fest €/h", "R4/R5 guardia especial €/h",
            "R1 con 80h neto mes", "R1 con 80h bruto mes", "R5 con 80h neto mes", "R5 con 80h bruto mes",
            "Horas ordinarias usadas 35h anualizadas", "Horas totales con 80h guardia",
            "R1 neto €/h real 35h+80h", "R5 neto €/h real 35h+80h",
            "R1 neto €/h si 37,5h+80h", "R5 neto €/h si 37,5h+80h",
        ])
        for ccaa, source, base, guard in ROWS:
            r1_net = WITH_GUARDS[ccaa][0][0]
            r5_net = WITH_GUARDS[ccaa][4][0]
            writer.writerow([
                display_ccaa(ccaa), source, *base, *guard[0], *guard[1], *guard[2], *guard[3],
                r1_net, WITH_GUARDS[ccaa][0][1], r5_net, WITH_GUARDS[ccaa][4][1],
                round(ORDINARY_HOURS, 1), round(TOTAL_HOURS, 1),
                round(r1_net / TOTAL_HOURS, 2), round(r5_net / TOTAL_HOURS, 2),
                round(r1_net / TOTAL_HOURS_37_5, 2), round(r5_net / TOTAL_HOURS_37_5, 2),
            ])
    return path


def write_xlsx() -> Path:
    path = OUT / "retribuciones_mir_2026.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Datos CCAA"
    headers = [
        "CCAA", "Fuente", "R1 sin guardias", "R2 sin guardias", "R3 sin guardias",
        "R4 sin guardias", "R5 sin guardias", "R1 lab", "R1 fest", "R1 especial",
        "R2 lab", "R2 fest", "R2 especial", "R3 lab", "R3 fest", "R3 especial",
        "R4/R5 lab", "R4/R5 fest", "R4/R5 especial", "R1 neto con 80h",
        "R1 bruto con 80h", "R5 neto con 80h", "R5 bruto con 80h",
        "Horas ordinarias 35h", "Horas totales 35h+80h", "R1 €/h 35h+80h",
        "R5 €/h 35h+80h", "R1 €/h 37,5h+80h", "R5 €/h 37,5h+80h",
    ]
    ws.append(headers)
    for ccaa, source, base, guard in ROWS:
        r1_net = WITH_GUARDS[ccaa][0][0]
        r5_net = WITH_GUARDS[ccaa][4][0]
        ws.append([
            display_ccaa(ccaa), source, *base, *guard[0], *guard[1], *guard[2], *guard[3],
            r1_net, WITH_GUARDS[ccaa][0][1], r5_net, WITH_GUARDS[ccaa][4][1],
            round(ORDINARY_HOURS, 1), round(TOTAL_HOURS, 1),
            round(r1_net / TOTAL_HOURS, 2), round(r5_net / TOTAL_HOURS, 2),
            round(r1_net / TOTAL_HOURS_37_5, 2), round(r5_net / TOTAL_HOURS_37_5, 2),
        ])

    header_fill = PatternFill("solid", fgColor="17324D")
    light_fill = PatternFill("solid", fgColor="EAF2F8")
    thin = Side(style="thin", color="D5DBDB")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.border = Border(bottom=thin)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
        if row[0].row % 2 == 0:
            for cell in row:
                cell.fill = light_fill
    for col in range(1, ws.max_column + 1):
        width = 16 if col > 2 else (22 if col == 1 else 24)
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "C2"
    ws.auto_filter.ref = ws.dimensions

    summary = wb.create_sheet("Resumen")
    summary.append(["Métrica", "R1", "R2", "R3", "R4", "R5"])
    no_guard = means_no_guards()
    wg = means_with_guards()
    summary.append(["Bruto medio sin guardias €/mes", *[round(x, 0) for x in no_guard]])
    summary.append(["Bruto medio cartel/80h €/mes", *[round(x[1], 0) for x in wg]])
    summary.append(["Neto medio cartel/80h €/mes", *[round(x[0], 0) for x in wg]])
    summary.append(["Guardias como % del bruto total", *[round((wg[i][1] - no_guard[i]) / wg[i][1], 3) for i in range(5)]])
    summary.append(["Horas ordinarias/mes usadas", ORDINARY_HOURS, ORDINARY_HOURS, ORDINARY_HOURS, ORDINARY_HOURS, ORDINARY_HOURS])
    summary.append(["Horas de guardia/mes usadas", GUARD_HOURS, GUARD_HOURS, GUARD_HOURS, GUARD_HOURS, GUARD_HOURS])
    summary.append(["Horas totales/mes", TOTAL_HOURS, TOTAL_HOURS, TOTAL_HOURS, TOTAL_HOURS, TOTAL_HOURS])
    summary.append(["Neto €/h real con 80h", *[round(wg[i][0] / TOTAL_HOURS, 2) for i in range(5)]])
    summary.append(["Neto €/h si 37,5 h/sem + 80h", *[round(wg[i][0] / TOTAL_HOURS_37_5, 2) for i in range(5)]])
    summary.append(["Neto €/h con 80h + 5h/sem formacion", *[round(wg[i][0] / TOTAL_HOURS_PLUS_FORMATION_5H, 2) for i in range(5)]])
    for cell in summary[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center")
    summary.column_dimensions["A"].width = 42
    for col in "BCDEF":
        summary.column_dimensions[col].width = 14
    chart = BarChart()
    chart.title = "Bruto medio: sin guardias vs cartel con 80h"
    chart.y_axis.title = "€/mes"
    chart.x_axis.title = "Año residencia"
    data = Reference(summary, min_col=2, max_col=6, min_row=2, max_row=3)
    cats = Reference(summary, min_col=2, max_col=6, min_row=1, max_row=1)
    chart.add_data(data, from_rows=True, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 7
    chart.width = 14
    summary.add_chart(chart, "A12")

    src = wb.create_sheet("Fuentes")
    src.append(["Fuente", "URL o archivo"])
    for name, url in SOURCES.items():
        src.append([name, url])
    src.column_dimensions["A"].width = 48
    src.column_dimensions["B"].width = 120
    for cell in src[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
    wb.save(path)
    return path


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], max_width: int, fnt, fill, line_spacing=6):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_spacing
    return y


def write_infographics() -> list[Path]:
    paths = []
    no_guard = means_no_guards()
    wg = means_with_guards()
    r1_guard_share = (wg[0][1] - no_guard[0]) / wg[0][1]
    r5_guard_share = (wg[4][1] - no_guard[4]) / wg[4][1]

    img = Image.new("RGB", (1600, 2000), "white")
    d = ImageDraw.Draw(img)
    navy = (23, 50, 77)
    red = (176, 53, 65)
    green = (40, 110, 86)
    gray = (82, 91, 103)
    pale = (239, 245, 248)
    d.rectangle([0, 0, 1600, 220], fill=navy)
    d.text((70, 52), "Lo que oculta el cartel sobre los MIR", font=font(56, True), fill="white")
    d.text((70, 130), "Los importes altos incluyen 80 horas mensuales de guardia: más de media jornada adicional.", font=font(30), fill=(220, 232, 238))

    cards = [
        ("Media R1", wg[0][0], wg[0][1], no_guard[0], r1_guard_share),
        ("Media R5", wg[4][0], wg[4][1], no_guard[4], r5_guard_share),
    ]
    x = 70
    for title, net, gross, base, share in cards:
        d.rounded_rectangle([x, 280, x + 700, 740], radius=22, fill=pale, outline=(205, 218, 226), width=2)
        d.text((x + 40, 320), title, font=font(36, True), fill=navy)
        d.text((x + 40, 385), f"{euro(net)} € netos/mes", font=font(48, True), fill=red)
        d.text((x + 40, 445), f"{euro(gross)} € brutos/mes", font=font(28, True), fill=gray)
        d.text((x + 40, 520), f"Sin guardias: {euro(base)} € brutos/mes", font=font(30), fill=navy)
        d.text((x + 40, 575), f"Guardias dentro del bruto: {share*100:.1f}%", font=font(34, True), fill=green)
        d.text((x + 40, 640), f"€/h neto real con 80 h: {euro(net / TOTAL_HOURS, 2)}", font=font(34, True), fill=red)
        x += 760

    d.text((70, 820), "La comparación correcta no es salario mensual contra salario mensual.", font=font(38, True), fill=navy)
    y = 890
    bullets = [
        f"Jornada ordinaria usada: {euro(ORDINARY_HOURS, 1)} h/mes ({ORDINARY_WEEKLY_HOURS} h/semana anualizadas).",
        "Guardias extra del escenario del cartel: 80 h/mes, tardes, noches, fines de semana y festivos.",
        f"Tiempo total mínimo: {euro(TOTAL_HOURS, 1)} h/mes, sin contar estudio, sesiones, cursos, investigación ni salidas tarde.",
        "El anteproyecto reconoce el problema: para residentes habla de revisar retribuciones y limitar guardias.",
    ]
    for b in bullets:
        d.ellipse([85, y + 10, 103, y + 28], fill=green)
        y = draw_wrapped(d, b, (125, y), 1350, font(31), gray, 9) + 20

    d.rectangle([70, 1230, 1530, 1800], fill=(250, 250, 250), outline=(215, 215, 215))
    d.text((105, 1275), "Mensaje corto para contestar al cartel", font=font(40, True), fill=navy)
    msg = (
        "No, un R1 no cobra 2.087 € netos por su jornada normal. Esa cifra sale de sumar unas 80 horas "
        "mensuales de guardias a un salario medio sin guardias de unos 1.447 € brutos. Si se divide por las "
        "231,7 horas realmente trabajadas, la media R1 queda alrededor de 9,0 € netos/hora, antes de contar el "
        "trabajo formativo no registrado."
    )
    draw_wrapped(d, msg, (105, 1350), 1380, font(34), (30, 37, 43), 11)
    d.text((70, 1870), "Fuentes: SIMEG 2026, Ministerio de Sanidad, Anteproyecto de Estatuto Marco 08/01/2026.", font=font(24), fill=gray)
    path = OUT / "infografia_cartel_vs_realidad_mir.png"
    img.save(path, quality=95)
    paths.append(path)

    img2 = Image.new("RGB", (1800, 1400), "white")
    d = ImageDraw.Draw(img2)
    d.rectangle([0, 0, 1800, 170], fill=navy)
    d.text((60, 45), "Retribución MIR 2026: el peso de las guardias por CCAA", font=font(50, True), fill="white")
    d.text((60, 105), "R1: bruto sin guardias frente a bruto con 80 h/mes de guardias extra.", font=font(26), fill=(220, 232, 238))
    sorted_rows = sorted(ROWS, key=lambda r: WITH_GUARDS[r[0]][0][1] - r[2][0], reverse=True)
    max_val = max(WITH_GUARDS[r[0]][0][1] for r in ROWS)
    y = 230
    for ccaa, source, base, guard in sorted_rows:
        with_g = WITH_GUARDS[ccaa][0][1]
        extra = with_g - base[0]
        d.text((60, y), display_ccaa(ccaa), font=font(24, True), fill=navy)
        x0 = 390
        bar_total = int((with_g / max_val) * 1050)
        bar_base = int((base[0] / max_val) * 1050)
        d.rectangle([x0, y + 2, x0 + bar_total, y + 30], fill=(224, 235, 240))
        d.rectangle([x0, y + 2, x0 + bar_base, y + 30], fill=green)
        d.rectangle([x0 + bar_base, y + 2, x0 + bar_total, y + 30], fill=red)
        d.text((1480, y - 2), f"{base[0]} + {extra} = {with_g} €", font=font(22), fill=gray)
        y += 60
    d.rectangle([390, 1270, 430, 1300], fill=green)
    d.text((445, 1265), "bruto sin guardias", font=font(24), fill=gray)
    d.rectangle([720, 1270, 760, 1300], fill=red)
    d.text((775, 1265), "guardias extra incluidas en el cartel", font=font(24), fill=gray)
    path2 = OUT / "infografia_ccaa_guardias_r1.png"
    img2.save(path2, quality=95)
    paths.append(path2)
    return paths


def make_markdown() -> Path:
    no_guard = means_no_guards()
    wg = means_with_guards()
    path = OUT / "dossier_mir_estatuto_marco.md"
    lines = []
    lines.append("# Dossier para rebatir el cartel de CSIF sobre sueldos MIR\n")
    lines.append("Fecha de trabajo: 6 de mayo de 2026. Datos retributivos principales: informe SIMEG 2026 con datos disponibles en abril de 2026, contrastado con tablas oficiales autonómicas localizadas.\n")
    lines.append("## Tesis central\n")
    lines.append("El cartel presenta como `sueldo mensual` unas cifras que incorporan guardias. Eso mezcla dos magnitudes distintas: la jornada ordinaria y una bolsa de trabajo adicional, penosa y en noches/festivos. Con la hipótesis usada por el propio estudio de referencia, esos importes incluyen 80 horas mensuales de guardia, más de media jornada ordinaria adicional si se toma como referencia una jornada de 35 h/semana.\n")
    lines.append("Nota metodológica: las medias del resumen usan la fila `Media de España` publicada en la tabla SIMEG/CTO; la tabla por CCAA separa la parte ordinaria y la parte atribuible a guardias. Para el cálculo por hora se usa 35 h/semana anualizadas: `35 x 52 / 12 = 151,7 h/mes`. Sumando 80 h de guardia, el divisor principal es `231,7 h/mes`. Si una CCAA aplica 37,5 h/semana o un cómputo anual superior, el divisor sube a `242,5 h/mes` y el salario neto por hora baja.\n")
    lines.append("Comentario metodológico adicional: no se han añadido horas estimadas de estudio ni formación al cálculo principal. Eso no significa que no existan. Muchos residentes dedican buena parte de sus tardes a estudiar, preparar sesiones, realizar postgrados o másteres, doctorado, asistir a congresos, investigar y publicar. Ese trabajo formativo suele no estar retribuido y, en no pocos casos, implica costes directos para el propio residente.\n")
    lines.append("## Cifras clave\n")
    lines.append("| Concepto | R1 media | R5 media |\n|---|---:|---:|")
    lines.append(f"| Bruto mensual sin guardias | {euro(no_guard[0])} € | {euro(no_guard[4])} € |")
    lines.append(f"| Neto mensual mostrado por el cartel/80 h | {euro(wg[0][0])} € | {euro(wg[4][0])} € |")
    lines.append(f"| Bruto mensual mostrado por el cartel/80 h | {euro(wg[0][1])} € | {euro(wg[4][1])} € |")
    lines.append(f"| Guardias dentro del bruto total | {(wg[0][1]-no_guard[0])/wg[0][1]*100:.1f}% | {(wg[4][1]-no_guard[4])/wg[4][1]*100:.1f}% |")
    lines.append(f"| Horas ordinarias usadas en el cálculo | {ORDINARY_HOURS:.1f} h | {ORDINARY_HOURS:.1f} h |")
    lines.append(f"| Guardias incluidas en el cartel | {GUARD_HOURS} h | {GUARD_HOURS} h |")
    lines.append(f"| Horas mensuales reales usadas en el cálculo | {TOTAL_HOURS:.1f} h | {TOTAL_HOURS:.1f} h |")
    lines.append(f"| Neto por hora real con 80 h de guardia | {wg[0][0]/TOTAL_HOURS:.2f} €/h | {wg[4][0]/TOTAL_HOURS:.2f} €/h |")
    lines.append(f"| Escenario 37.5 h/semana, si aplica | {wg[0][0]/TOTAL_HOURS_37_5:.2f} €/h | {wg[4][0]/TOTAL_HOURS_37_5:.2f} €/h |\n")
    lines.append("## Qué dice realmente el anteproyecto\n")
    lines.append("- El Ministerio presenta el texto de enero de 2026 como último borrador negociado, orientado a estabilidad, conciliación, prevención de riesgos y regulación de jornadas.")
    lines.append("- El anteproyecto fija una duración máxima conjunta de jornada ordinaria, guardia y otras modalidades de 45 h semanales de promedio cuatrimestral.")
    lines.append("- La guardia del personal médico/facultativo se regula como horario adicional para urgencia. En laborables se limita a 17 h efectivas; excepcionalmente puede llegar a 24 h con consentimiento y requisitos preventivos.")
    lines.append("- Para personal sanitario en formación especializada, la disposición adicional vigesimosegunda plantea revisar el marco retributivo, limitar con carácter general a 80 h mensuales y a cuatro módulos de guardia física al mes salvo excepciones.")
    lines.append("- El propio texto, por tanto, reconoce que la guardia del residente es una carga específica que no puede venderse como sueldo ordinario.\n")
    lines.append("## Situación sindical y profesional\n")
    lines.append("- SATSE-FSES, FSS-CCOO, UGT y CSIF comunicaron acuerdo con Sanidad para iniciar la tramitación del anteproyecto.")
    lines.append("- El Comité de Huelga formado por CESM, SMA, Metges de Catalunya, AMYTS, Sindicato Médico de Euskadi y O'MEGA mantiene oposición y movilizaciones, reclamando un estatuto/ámbito propio médico-facultativo.")
    lines.append("- Los puntos de choque son: regulación propia de jornada y guardias, cómputo para jubilación y descansos, clasificación profesional, retribuciones, incompatibilidades/dedicación exclusiva y movilidad.\n")
    lines.append("## Tabla comparativa nacional R1\n")
    lines.append("| CCAA | R1 bruto sin guardias | R1 bruto con 80 h | Parte atribuible a guardias | R1 guardia laborable €/h | R1 festivo €/h |")
    lines.append("|---|---:|---:|---:|---:|---:|")
    for ccaa, source, base, guard in ROWS:
        with_g = WITH_GUARDS[ccaa][0][1]
        lines.append(f"| {display_ccaa(ccaa)} | {base[0]} € | {with_g} € | {with_g-base[0]} € | {guard[0][0]:.2f} | {guard[0][1]:.2f} |")
    lines.append("\n## Respuesta breve para redes o asamblea\n")
    lines.append("> El cartel de CSIF no muestra el sueldo normal de un MIR: suma unas 80 horas mensuales de guardias. Un R1 medio no pasa de 1.447 € brutos sin guardias; la cifra de 2.087 € netos sale de trabajar unas 231,7 h/mes si usamos 35 h/semana anualizadas. Eso deja la media R1 en torno a 9,0 € netos/h antes de contar estudio, sesiones, postgrados, congresos, investigación, publicaciones y salidas tarde. Si el cómputo ordinario aplicable es de 37,5 h/semana, baja a 8,6 € netos/h.\n")
    lines.append("## Fuentes\n")
    for name, url in SOURCES.items():
        lines.append(f"- {name}: {url}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def make_argumentario() -> Path:
    wg = means_with_guards()
    no_guard = means_no_guards()
    r1_hour = wg[0][0] / TOTAL_HOURS
    r1_hour_37_5 = wg[0][0] / TOTAL_HOURS_37_5
    r1_guard_share = (wg[0][1] - no_guard[0]) / wg[0][1]
    path = OUT / "argumentario_rapido_csif_mir.md"
    lines = [
        "# Argumentario rápido contra el cartel de CSIF",
        "",
        "## Mensaje principal",
        "",
        f"El cartel no muestra el sueldo normal de un MIR: mezcla sueldo ordinario con unas {GUARD_HOURS} horas mensuales de guardias. Un R1 medio está en torno a {euro(no_guard[0])} euros brutos mensuales sin guardias; la cifra de {euro(wg[0][0])} euros netos aparece al sumar trabajo adicional en tardes, noches, fines de semana y festivos.",
        "",
        "## Respuesta de 30 segundos",
        "",
        f"No es correcto decir que “los MIR cobran 2.000-3.000 euros” sin explicar las horas. Esas cifras incluyen unas {GUARD_HOURS} horas de guardia al mes. Con jornada ordinaria anualizada de {ORDINARY_WEEKLY_HOURS} h/semana, eso son unas {euro(TOTAL_HOURS, 1)} h/mes trabajadas. En un R1 medio, {euro(wg[0][0])} euros netos entre esas horas son unos {euro(r1_hour, 1)} euros netos/hora, antes de contar estudio, sesiones, investigación, cursos y salidas tarde. Si el cómputo aplicable es 37,5 h/semana, el divisor sube a {euro(TOTAL_HOURS_37_5, 1)} h/mes y el R1 baja a {euro(r1_hour_37_5, 1)} euros netos/hora.",
        "",
        "## Lo que ni siquiera entra en el cálculo",
        "",
        "No se han sumado horas estimadas de estudio ni formación para inflar el denominador. Pero existen: muchos residentes pasan buena parte de las tardes estudiando, preparando sesiones, haciendo postgrados o másteres, doctorado, congresos, investigación y publicaciones. Ese trabajo no suele estar retribuido y, en bastantes casos, implica costes directos para el propio residente.",
        "",
        "## Tres preguntas para desmontar la tabla",
        "",
        "1. ¿Cuánto es sueldo ordinario sin guardias?",
        "2. ¿Cuántas horas reales se han trabajado para llegar a ese neto?",
        "3. ¿Por qué se compara un salario con guardias, noches y festivos con salarios ordinarios de otras profesiones?",
        "",
        "## Datos cortos",
        "",
        "| Dato | Cifra |",
        "|---|---:|",
        f"| R1 medio bruto sin guardias | {euro(no_guard[0])} €/mes |",
        f"| R1 medio neto con 80 h de guardia | {euro(wg[0][0])} €/mes |",
        f"| Horas ordinarias/mes usadas | {euro(ORDINARY_HOURS, 1)} h |",
        f"| Guardias añadidas | {GUARD_HOURS} h |",
        f"| Total mínimo de horas/mes | {euro(TOTAL_HOURS, 1)} h |",
        f"| Neto real R1 con 80 h | {euro(r1_hour, 2)} €/h |",
        f"| Neto R1 si aplica 37,5 h/semana | {euro(r1_hour_37_5, 2)} €/h |",
        f"| Peso de guardias en bruto R1 | {euro(r1_guard_share * 100, 1)}% |",
        "",
        "## Frase final",
        "",
        "Si casi la mitad del bruto depende de guardias, no estamos hablando de “cobrar mucho”: estamos hablando de sostener el sistema con jornadas extendidas, noches, festivos y responsabilidad clínica creciente.",
        "",
        "Fuentes: SIMEG 2026, Ministerio de Sanidad, Anteproyecto de Estatuto Marco 08/01/2026 y tablas oficiales autonómicas localizadas.",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def make_docx(md_path: Path, images: list[Path]) -> Path:
    no_guard = means_no_guards()
    wg = means_with_guards()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.color.rgb = RGBColor(23, 50, 77)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.color.rgb = RGBColor(23, 50, 77)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Dossier para rebatir el cartel de CSIF sobre sueldos MIR")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(23, 50, 77)
    sub = doc.add_paragraph("Fecha: 6 de mayo de 2026. Documento de trabajo con datos contrastados y cálculos reproducibles.")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Resumen ejecutivo", level=1)
    for text in [
        "El cartel presenta como sueldo mensual cifras que incluyen guardias. No es una comparación válida con una jornada ordinaria.",
        "La hipotesis de referencia son 80 horas mensuales de guardia que se suman a la jornada ordinaria.",
        "Con 35 h/semana anualizadas, el R1 medio mostrado como 2.087 euros netos/mes equivale a unas 231,7 horas mensuales y queda en torno a 9,0 euros netos/hora.",
        "Si una comunidad aplica 37,5 h/semana o un cómputo anual superior, el divisor sube a 242,5 horas mensuales y el R1 baja a unos 8,6 euros netos/hora.",
        "El anteproyecto reconoce el problema al regular límites de guardias y pedir revisión retributiva para el personal en formación sanitaria especializada.",
    ]:
        doc.add_paragraph(text, style=None).style = styles["Normal"]

    doc.add_heading("Cifras clave", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Concepto"
    hdr[1].text = "R1 media"
    hdr[2].text = "R5 media"
    rows = [
        ("Bruto mensual sin guardias", f"{euro(no_guard[0])} €", f"{euro(no_guard[4])} €"),
        ("Neto mensual cartel/80 h", f"{euro(wg[0][0])} €", f"{euro(wg[4][0])} €"),
        ("Bruto mensual cartel/80 h", f"{euro(wg[0][1])} €", f"{euro(wg[4][1])} €"),
        ("Guardias dentro del bruto total", f"{(wg[0][1]-no_guard[0])/wg[0][1]*100:.1f}%", f"{(wg[4][1]-no_guard[4])/wg[4][1]*100:.1f}%"),
        ("Horas ordinarias 35 h/semana", f"{ORDINARY_HOURS:.1f} h", f"{ORDINARY_HOURS:.1f} h"),
        ("Guardias incluidas en el cartel", f"{GUARD_HOURS} h", f"{GUARD_HOURS} h"),
        ("Horas mensuales reales", f"{TOTAL_HOURS:.1f} h", f"{TOTAL_HOURS:.1f} h"),
        ("Neto por hora real", f"{wg[0][0]/TOTAL_HOURS:.2f} €/h", f"{wg[4][0]/TOTAL_HOURS:.2f} €/h"),
        ("Escenario 37,5 h/semana", f"{wg[0][0]/TOTAL_HOURS_37_5:.2f} €/h", f"{wg[4][0]/TOTAL_HOURS_37_5:.2f} €/h"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(parse_shading("17324D"))

    doc.add_heading("Infografia principal", level=1)
    doc.add_picture(str(images[0]), width=Inches(6.8))

    doc.add_heading("Puntos para argumentar", level=1)
    points = [
        "Pedir que cualquier tabla separe sueldo ordinario, guardias, complementos y neto estimado.",
        "Convertir siempre a euros/hora con horas reales: jornada ordinaria mensual + guardias + tiempo formativo no registrado.",
        "Recordar que las guardias son trabajo adicional en tardes/noches/festivos y no una mejora salarial comparable con una nómina ordinaria.",
        "Citar el propio anteproyecto: si limita guardias en residentes y exige revisar retribuciones, no puede venderse la guardia como salario normal.",
    ]
    for p in points:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("Fuentes", level=1)
    for name, url in SOURCES.items():
        doc.add_paragraph(f"{name}: {url}", style=None)

    path = OUT / "dossier_mir_estatuto_marco.docx"
    doc.save(path)
    return path


def parse_shading(fill: str):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    return parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>')


def main() -> None:
    csv_path = write_csv()
    xlsx_path = write_xlsx()
    images = write_infographics()
    md_path = make_markdown()
    argumentario_path = make_argumentario()
    docx_path = make_docx(md_path, images)
    print("created")
    for p in [csv_path, xlsx_path, *images, md_path, argumentario_path, docx_path]:
        print(p)


if __name__ == "__main__":
    main()
